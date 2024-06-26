from nonebot import require

require("nonebot_plugin_localstore")

import nonebot_plugin_localstore as store
import asyncio
import os
from nonebot.log import logger
from .openai_func import request_one
from .config import  config
from docx import Document # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
                


async def send_request(content):

    message = f"""Generate part of the paper according to the following outline.
    =====
    {content}
    ===== 
    Note: Only respond the content of the paper.
"""
    messages=[
            {"role": "user", "content": message}
        ]
    data={
        "model": config.model,
        "messages": messages,
        "stream": False,
    } 
        
    try:
        response =await request_one(data)
        if response == None:
                response =await request_one(data)



    except Exception as e:
            return str(e)
    if response=="error":
        return "=====尝试生成该部分内容时发生网络错误请稍后再试或尝试切换模型====="  

    response=response['choices'][0]['message']['content']
    return response

async def generate_paper(file_path, topic, user_id: str,parts):




    messages=[{"role": "user", "content":f"写一个要求或主题为：\n{topic}\n的论文提纲由{parts}部分组成,每个部分之间用=====加换行隔开，仅需要回复提纲内容，不能回复其他"}]
    data={
        "model": config.model,
        "messages": messages,
        "stream": False,
    } 
    try:
        logger.info("第一次尝试生成提纲")
        response =await request_one(data)
        if response == None:
                response =await request_one(data)

    except Exception as e:
        logger.info(e)
        try:
            logger.info("第二次尝试生成提纲")
            response =await request_one(data)                     
        except Exception as e:
            logger.info(e)

    if response=="error":
        return "生成提纲时发生错误，请稍后再试或尝试切换模型"


    
    aaa=response['choices'][0]['message']['content']
    logger.info("论文提纲为"+aaa)
    contents=aaa.split("=====")

    tasks = [send_request(content) for content in contents]
    try:
        results = await asyncio.gather(*tasks)
    except Exception as a:
        logger.info(a)
    res = "".join(results)
    res  = res .replace('#', '')
    res  = res .replace('*', '')
    res = res.replace('`','')
    def create_word_document(text, file_path):
        # 创建一个新的Word文档
        document = Document()
        
        # 按段落添加文本
        paragraphs = text.split('\n')
        
        for i, para in enumerate(paragraphs):
            paragraph = document.add_paragraph(para)
            if i == 0:  # 如果是第一行
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 1:  # 如果是第二行
                author_paragraph = document.add_paragraph("作者：")
                author_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 保存文档

        document.save(file_path)


    folder = store.get_cache_dir("nonebot_plugin_multigpt") / user_id
    folder.mkdir(parents=True, exist_ok=True)

    # 定义文件路径
    file_path = folder / f"{user_id}.docx"

    # 创建 Word 文档
    create_word_document(res, file_path)

    # 删除临时文件
    dir_path = folder
    prefix = "prefix_"

    for file_name in os.listdir(dir_path):
        if file_name.startswith(prefix):
            file_path = dir_path / file_name
            if file_path.is_file():
                os.remove(file_path)

    # 返回当前工作目录和文件路径
    return str(store.get_cache_dir("nonebot_plugin_multigpt") / user_id / f"{user_id}.docx")